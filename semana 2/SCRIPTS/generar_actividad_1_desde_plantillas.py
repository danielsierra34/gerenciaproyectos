from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "INPUT"
OUTPUT = ROOT / "OUTPUT"
BASE_DATE = "12 de abril de 2026"

objectives = [
    ("Implementar funcionalidades de compraventa de inmuebles en línea", "Tener en operación, antes de finalizar el primer año de concesión, el 100% de las funcionalidades priorizadas: cita virtual, biometría, gestión documental, anotaciones y flujo de trabajo."),
    ("Reducir la necesidad de presencialidad en notarías", "Permitir que al menos el 80% de las etapas de lectura, revisión y comparecencia del trámite de compraventa se realicen virtualmente para los usuarios habilitados."),
    ("Mantener seguridad de identidad equivalente o superior al proceso presencial", "Validar identidad mediante al menos tres factores biométricos: facial, voz y huella; con registro de auditoría para el 100% de las sesiones notariales virtuales."),
    ("Asegurar desempeño y disponibilidad de la solución", "Definir y cumplir tiempos de respuesta máximos para operaciones críticas y una disponibilidad objetivo mínima de 99,5% en horario de servicio durante el piloto en ciudad capital."),
    ("Evitar sanciones contractuales por atraso", "Cerrar desarrollo, pruebas y puesta en producción dentro de los 12 meses, evitando multas del 10% del valor total por cada mes o fracción de retraso."),
]

scope_in = [
    "Análisis de infraestructura de las entidades participantes: conectividad, servidores, software, seguridad, ambientes y protocolos de comunicación.",
    "Levantamiento de demanda de operaciones frente a capacidad instalada actual, futura y de mediano plazo.",
    "Diseño e implementación de programación de citas virtuales para las partes de la compraventa.",
    "Integración de mecanismos de identificación facial, por voz y por huella.",
    "Gestión de escrituras y contenidos documentales en formato electrónico.",
    "Funcionalidad para anotaciones sobre documentos por parte de participantes y revisión en notaría.",
    "Automatización del flujo de trabajo asociado a compraventa de inmuebles.",
    "Propuestas técnicas de escalabilidad, desempeño, disponibilidad y seguridad.",
    "Pruebas funcionales, técnicas, de seguridad e integración antes de la salida a producción.",
    "Puesta en operación de las funcionalidades al cierre del primer año de concesión.",
]

scope_out = [
    "Implementación nacional completa fuera de la ciudad capital durante esta fase.",
    "Digitalización de todos los trámites notariales distintos a compraventa de inmuebles.",
    "Compra o modernización completa de infraestructura física de todas las entidades si supera el presupuesto aprobado.",
    "Reforma normativa o definición de políticas públicas ajenas al alcance técnico del proyecto.",
    "Operación permanente de soporte posterior a la estabilización inicial, salvo transferencia y garantía acordadas.",
]

assumptions = [
    ("La SGNP entregará información histórica de demanda y planes de crecimiento.", "Validar con acta de entrega de información y revisión de completitud durante el inicio del proyecto."),
    ("Tinterillos mantendrá la alianza tecnológica con Turing durante el primer año de concesión.", "Validar con contrato/orden de trabajo firmado y comité de seguimiento mensual."),
    ("El aplicativo actual basado en microservicios Spring Boot permite integrar nuevos servicios sin rediseñar toda la plataforma.", "Validar mediante evaluación de arquitectura y prueba técnica de integración temprana."),
    ("Las notarías de ciudad capital contarán con conectividad y ambientes mínimos para pruebas y operación.", "Validar con diagnóstico de infraestructura por entidad y plan de cierre de brechas."),
    ("Los usuarios aceptarán mecanismos biométricos para identidad digital en el trámite.", "Validar con lineamientos legales, pruebas piloto y aceptación de usuarios representativos."),
]

restrictions = [
    "Presupuesto de sistematización financiado por SGNP: USD 1.000.000.",
    "Fecha límite: funcionalidades desarrolladas, probadas y en operación al finalizar el primer año de concesión.",
    "Multa contractual estimada: 10% del valor total del proyecto por cada mes o fracción de atraso.",
    "Disponibilidad limitada de recursos de Tinterillos por ejecución simultánea de varios proyectos.",
    "La solución debe integrarse con el aplicativo actual basado principalmente en microservicios Spring Boot.",
    "El proyecto se limita inicialmente a las notarías de la ciudad capital.",
    "La solución no puede reducir la seguridad actual sobre identidad de las partes frente al proceso presencial.",
    "El apoyo político tiene ventana de dos años del gobierno actual; no hay certeza sobre continuidad futura.",
]

risks = [
    ("Retraso en desarrollo o pruebas de funcionalidades críticas", "Alta", "Multas del 10% por mes o fracción, deterioro reputacional de Turing y riesgo para la licitación de Tinterillos.", "Gerente del proyecto Turing"),
    ("Infraestructura de notarías o entidades participantes no idónea", "Media", "Bajo desempeño, fallas de integración o imposibilidad de operar citas virtuales.", "Arquitecto de software"),
    ("Fallas en integración biométrica facial, voz y huella", "Media", "Riesgo de suplantación, rechazo legal o necesidad de mantener presencialidad.", "Líder técnico de seguridad"),
    ("Información histórica incompleta para dimensionamiento de demanda", "Media", "Subdimensionamiento de capacidad, costos no previstos y degradación del servicio.", "Analista de capacidad"),
    ("Resistencia de notarios, funcionarios o ciudadanos al trámite virtual", "Media", "Baja adopción, solicitudes de cambios y presión para mantener procesos manuales.", "Líder de gestión del cambio"),
    ("Cambios de alcance por nuevas funcionalidades o solicitudes políticas", "Alta", "Aumento de costo y tiempo, pérdida de foco en compraventa de inmuebles.", "Comité de control de cambios"),
    ("Vulnerabilidades de confidencialidad o integridad documental", "Media", "Exposición de escrituras, pérdida de confianza y posibles sanciones.", "Oficial de seguridad"),
    ("Dependencia de recursos clave de Turing", "Media", "Cuellos de botella en arquitectura, desarrollo o pruebas.", "Gerente de operaciones Turing"),
    ("Falta de continuidad política tras el gobierno actual", "Baja", "Reducción de apoyo para escalamiento o fases posteriores.", "Patrocinador SGNP"),
    ("Tiempos de respuesta insuficientes en operaciones críticas", "Media", "Mala experiencia de usuario, incumplimiento de criterios de desempeño y reprocesos.", "Líder de infraestructura"),
]

stakeholders = [
    ("1", "SGNP", "Patrocinador y entidad financiadora del proyecto."),
    ("2", "Tinterillos SA", "Contratista principal de la licitación y responsable ante SGNP."),
    ("3", "Turing", "Aliado tecnológico encargado de ejecutar el proyecto."),
    ("4", "Alcaldesa de ciudad capital", "Impulsora política de la sistematización y reducción de tramitología."),
    ("5", "Notarías de ciudad capital", "Entidades operadoras del trámite y usuarias institucionales."),
    ("6", "Notarios", "Responsables de la comparecencia y validez del trámite."),
    ("7", "Ciudadanos compradores y vendedores", "Usuarios finales del trámite de compraventa de inmuebles."),
    ("8", "Gerente general de Turing", "Promotor de la estrategia de incursión en proyectos con el Estado."),
    ("9", "Gerente financiero de Turing", "Evalúa rentabilidad, costos, multas y viabilidad económica."),
    ("10", "Gerente de operaciones de Turing", "Responsable de disponibilidad de recursos y capacidad de ejecución."),
    ("11", "Arquitecto de software de Turing", "Responsable técnico de arquitectura, integración y escalabilidad."),
    ("12", "Equipo técnico de Turing", "Analistas, desarrolladores, QA, DevOps y seguridad que construyen la solución."),
    ("13", "Comité evaluador de Mejor Iniciativa del Año", "Evalúa propuesta interna y reconocimiento de los gerentes jóvenes."),
    ("14", "Futuros gobiernos", "Podrían apoyar o frenar continuidad y escalamiento posterior."),
    ("15", "Entidades externas de identidad/biometría", "Proveedores o integradores de servicios de verificación de identidad."),
]

milestones = [
    ("Inicio formal del proyecto", "Abril de 2026"),
    ("Diagnóstico de infraestructura y demanda completado", "Mayo de 2026"),
    ("Arquitectura objetivo y plan de integración aprobados", "Junio de 2026"),
    ("Prototipo de cita virtual, gestión documental y anotaciones", "Agosto de 2026"),
    ("Integración biométrica facial, voz y huella completada", "Octubre de 2026"),
    ("Pruebas integrales, seguridad, desempeño y disponibilidad", "Diciembre de 2026 - febrero de 2027"),
    ("Piloto controlado en notarías seleccionadas", "Febrero de 2027"),
    ("Salida a producción en notarías de ciudad capital", "Marzo de 2027"),
    ("Cierre del primer año de concesión con funcionalidades en operación", "Abril de 2027"),
]

power_interest = [
    ("SGNP", "Alto", "Alto", "Gestionar de cerca"),
    ("Tinterillos SA", "Alto", "Alto", "Gestionar de cerca"),
    ("Turing", "Alto", "Alto", "Gestionar de cerca"),
    ("Alcaldesa de ciudad capital", "Alto", "Alto", "Gestionar de cerca"),
    ("Notarías de ciudad capital", "Alto", "Alto", "Gestionar de cerca"),
    ("Notarios", "Alto", "Alto", "Gestionar de cerca"),
    ("Ciudadanos compradores y vendedores", "Bajo", "Alto", "Mantener informados"),
    ("Gerente general de Turing", "Alto", "Alto", "Gestionar de cerca"),
    ("Gerente financiero de Turing", "Alto", "Medio", "Mantener satisfecho"),
    ("Gerente de operaciones de Turing", "Alto", "Alto", "Gestionar de cerca"),
    ("Arquitecto de software de Turing", "Medio", "Alto", "Mantener informados e involucrados"),
    ("Equipo técnico de Turing", "Medio", "Alto", "Mantener informados e involucrados"),
    ("Comité evaluador de Mejor Iniciativa del Año", "Medio", "Medio", "Monitorear y reportar avances clave"),
    ("Futuros gobiernos", "Alto", "Bajo", "Mantener satisfechos"),
    ("Entidades externas de identidad/biometría", "Medio", "Medio", "Gestionar por contrato y acuerdos de servicio"),
]

stakeholder_management = [
    ("SGNP", "Cumplir el contrato, lograr sistematización segura y usar adecuadamente el presupuesto.", "Comité directivo quincenal, reportes de avance, riesgos y control de cambios.", "Puede exigir alcance adicional; solución: priorización formal y aprobación de cambios.", "Financiación; legitimidad institucional; datos históricos; capacidad de convocatoria.", "Tiempo de comités; aprobaciones; seguimiento contractual; gestión de cambios.", "Atraso por aprobaciones; alternativa: matriz RACI y tiempos máximos de decisión."),
    ("Tinterillos SA", "Cumplir la licitación y proteger su reputación como contratista principal.", "Gestión de cuenta, tablero ejecutivo y acuerdos claros de responsabilidad con Turing.", "Puede trasladar presión contractual a Turing; solución: acuerdos de servicio y escalamiento.", "Acceso al cliente; conocimiento del contrato; coordinación con SGNP; soporte comercial.", "Coordinación contractual; reuniones; gestión de dependencias; atención de auditorías.", "Disponibilidad limitada de recursos; alternativa: plan de suplencias y decisiones priorizadas."),
    ("Turing", "Ejecutar con éxito, entrar en proyectos del Estado y evitar multas.", "Gobernanza interna semanal, control de alcance y gestión temprana de riesgos técnicos.", "La competencia interna puede sesgar prioridades; solución: alinear concurso con objetivos contractuales.", "Capacidad técnica; oportunidad estratégica; aprendizaje; posible crecimiento comercial.", "Asignación de equipo; costo de desarrollo; pruebas; gestión de proveedores.", "Sobrecarga del equipo; alternativa: plan de capacidad, contratación temporal y reservas de gestión."),
    ("Notarías y notarios", "Mantener validez, seguridad y eficiencia del trámite notarial.", "Mesas de levantamiento, pilotos, capacitación y canal de soporte.", "Pueden resistir cambios al proceso; solución: pruebas controladas y evidencia de seguridad.", "Conocimiento operativo; adopción; retroalimentación; legitimidad del proceso.", "Capacitación; tiempo de usuarios expertos; ajustes de procedimiento; soporte inicial.", "Baja adopción; alternativa: embajadores por notaría y plan de gestión del cambio."),
    ("Ciudadanos compradores y vendedores", "Realizar trámites con menos desplazamientos, seguridad y tiempos razonables.", "Comunicación clara, pruebas de usabilidad, soporte y mecanismos de contingencia.", "Desconfianza sobre biometría o documentos electrónicos; solución: pedagogía y consentimiento informado.", "Mayor adopción; retroalimentación; reducción de filas; validación de experiencia.", "Soporte; comunicación; atención de quejas; ajustes de accesibilidad.", "Rechazo del canal digital; alternativa: piloto gradual y canal asistido."),
    ("Gerencias y arquitectura de Turing", "Controlar rentabilidad, recursos, calidad técnica y visibilidad interna.", "Reportes internos, revisiones de arquitectura y control financiero mensual.", "Conflicto entre costo y calidad; solución: criterios de aceptación y análisis de impacto.", "Dirección técnica; priorización; control de costos; apoyo ejecutivo.", "Tiempo directivo; revisiones; ajustes de plan; gestión de escalamiento.", "Decisiones tardías; alternativa: comité de decisiones con agenda y responsables."),
]

benefit_cost_rows = [
    ("Beneficio 1", "Disminución de desplazamientos presenciales para ciudadanos en compraventa de inmuebles."),
    ("Beneficio 2", "Mejor trazabilidad del trámite mediante documentos electrónicos, anotaciones y auditoría."),
    ("Beneficio 3", "Posicionamiento de Turing en proyectos del Estado y posibles contratos futuros."),
    ("Beneficio 4", "Mayor eficiencia operativa para notarías y SGNP al automatizar flujos de trabajo."),
    ("Costo 1", "Tiempo de participación de SGNP, Tinterillos, notarías y equipo técnico en comités y levantamientos."),
    ("Costo 2", "Inversión en desarrollo, pruebas, seguridad, integraciones biométricas e infraestructura requerida."),
    ("Costo 3", "Costo de gestión del cambio, capacitación, soporte y comunicación a usuarios."),
    ("Costo 4", "Costo de control de riesgos, auditorías, pruebas de desempeño y contingencias por retrasos."),
]

risk_alternatives = [
    ("Riesgo de atraso", "Plan de entregas incrementales, hitos mensuales, ruta crítica y escalamiento temprano."),
    ("Riesgo de seguridad e identidad", "Pruebas de seguridad, validación legal, autenticación multifactor y auditoría de sesiones."),
    ("Riesgo de infraestructura insuficiente", "Diagnóstico temprano, dimensionamiento, pruebas de carga y plan de cierre de brechas."),
    ("Riesgo de resistencia al cambio", "Pilotos, capacitación, usuarios líderes por notaría y canal de soporte."),
    ("Riesgo de cambios de alcance", "Comité de control de cambios, backlog priorizado y evaluación de impacto en costo/tiempo."),
]

conflicts = [
    ("SGNP/Tinterillos vs. Turing", "Presión por nuevas funcionalidades frente a plazo fijo; resolver con control formal de cambios."),
    ("Notarios vs. automatización", "Temor a pérdida de control del acto notarial; resolver con validaciones, capacitación y participación en pilotos."),
    ("Ciudadanos vs. biometría", "Preocupaciones por privacidad; resolver con consentimiento, transparencia y controles de seguridad."),
    ("Gerencia financiera vs. equipo técnico", "Control de costos frente a necesidades de calidad; resolver con criterios de aceptación y priorización por riesgo."),
    ("Alcaldía/SGNP vs. futuros gobiernos", "Continuidad del apoyo político; resolver documentando beneficios y resultados medibles del primer año."),
]


def set_cell(cell, text):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(str(text))


def normalize_alignment(document):
    for paragraph in document.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def clear_document_metadata(document):
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = ""
    document.core_properties.keywords = ""
    document.core_properties.subject = ""


def bullet_lines(items):
    return "\n".join(f"- {item}" for item in items)


def objective_lines():
    return "\n\n".join(f"{objective}\nMétrica: {metric}" for objective, metric in objectives)


def assumption_lines():
    return "\n\n".join(f"- {assumption}\n  Cómo se valida: {validation}" for assumption, validation in assumptions)


def risk_lines():
    return "\n\n".join(
        f"- {risk}\n  Probabilidad: {probability}. Consecuencia: {impact.rstrip('.')}. Dueño: {owner}."
        for risk, probability, impact, owner in risks
    )


def stakeholder_lines():
    return "\n".join(f"{number}. {name}: {description}" for number, name, description in stakeholders)


def milestone_lines():
    return "\n".join(f"- {milestone}: {date}" for milestone, date in milestones)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table_before(document, before_paragraph, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header)
        for paragraph in table.rows[0].cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell(cells[index], value)

    before_paragraph._p.addprevious(table._tbl)
    return table


def add_paragraph_before(document, before_paragraph, text="", style=None):
    paragraph = document.add_paragraph(text, style=style)
    before_paragraph._p.addprevious(paragraph._p)
    return paragraph


def find_copyright_paragraph(document):
    for paragraph in document.paragraphs:
        if "Derechos Reservados" in paragraph.text:
            return paragraph
    return document.paragraphs[-1]


def build_charter_from_template():
    document = Document(INPUT / "Acta-de-Constitucion-.docx")

    header_table = document.tables[1]
    set_cell(header_table.cell(1, 1), "Aplicativo en línea para trámites notariales de compraventa de inmuebles de la SGNP")
    set_cell(header_table.cell(2, 1), "Superintendencia Gubernamental de Notario Público (SGNP)")
    set_cell(header_table.cell(2, 3), "Gerente de proyecto designado por Turing")
    set_cell(header_table.cell(3, 1), "Tinterillos SA / SGNP")
    set_cell(header_table.cell(3, 3), "Turing como aliado tecnológico de Tinterillos SA")
    set_cell(header_table.cell(4, 1), BASE_DATE)
    set_cell(header_table.cell(4, 3), "Abril de 2027")

    content_table = document.tables[2]
    problem = (
        "Las notarías de la ciudad capital dependen de trámites presenciales para procesos de alta demanda como la compraventa de inmuebles. "
        "Esto obliga a compradores y vendedores a desplazarse para leer escrituras y comparecer ante notario, lo que aumenta tiempos, tramitología y costos para los ciudadanos. "
        "El problema se alinea con la planeación estratégica de la alcaldía y de la SGNP, orientada a sistematizar las notarías, reducir trámites presenciales y facilitar la vida de los ciudadanos. "
        "Para resolverlo se desarrollarán funcionalidades en línea sobre el aplicativo unificado: cita virtual, identificación biométrica, manejo de escrituras electrónicas, anotaciones documentales y automatización del flujo de compraventa, manteniendo la seguridad de identidad del proceso actual."
    )
    set_cell(content_table.cell(2, 0), problem)
    set_cell(content_table.cell(2, 1), problem)
    set_cell(content_table.cell(5, 0), objective_lines())
    set_cell(content_table.cell(5, 1), "La evaluación se hará contra estos criterios: 100% de funcionalidades priorizadas, 80% de etapas virtualizables, 3 factores biométricos, 100% de sesiones auditadas, disponibilidad mínima de 99,5% y entrega dentro de 12 meses.")
    set_cell(content_table.cell(8, 0), bullet_lines(scope_in))
    set_cell(content_table.cell(8, 1), bullet_lines(scope_out))
    set_cell(content_table.cell(11, 0), assumption_lines())
    set_cell(content_table.cell(11, 1), assumption_lines())
    set_cell(content_table.cell(14, 0), bullet_lines(restrictions))
    set_cell(content_table.cell(14, 1), bullet_lines(restrictions))
    set_cell(content_table.cell(17, 0), risk_lines())
    set_cell(content_table.cell(17, 1), risk_lines())
    set_cell(content_table.cell(20, 0), stakeholder_lines())
    set_cell(content_table.cell(20, 1), stakeholder_lines())
    set_cell(content_table.cell(23, 0), milestone_lines())
    set_cell(content_table.cell(23, 1), milestone_lines())

    normalize_alignment(document)
    clear_document_metadata(document)
    OUTPUT.mkdir(exist_ok=True)
    document.save(OUTPUT / "Acta-de-Constitucion-diligenciada.docx")


def build_matrix_from_template():
    document = Document(INPUT / "Matriz_poder_interes.docx")

    stakeholder_table = document.tables[2]
    for index, (_, name, description) in enumerate(stakeholders, start=1):
        set_cell(stakeholder_table.cell(index, 1), f"{name}: {description}")

    before = find_copyright_paragraph(document)
    add_paragraph_before(document, before, "Clasificación en la Matriz de Poder - Interés", style="Heading 1")
    add_table_before(document, before, ["Interesado", "Poder", "Interés", "Estrategia"], power_interest)

    add_paragraph_before(document, before, "Cuadrantes de gestión", style="Heading 1")
    add_table_before(
        document,
        before,
        ["Cuadrante", "Interesados", "Gestión sugerida"],
        [
            ("Alto poder / Alto interés", "SGNP, Tinterillos, Turing, Alcaldesa, Notarías, Notarios, Gerente general y Gerente de operaciones", "Gestionar de cerca mediante comités, reportes ejecutivos, decisiones oportunas y control de cambios."),
            ("Alto poder / Bajo o medio interés", "Gerente financiero de Turing y futuros gobiernos", "Mantener satisfechos con indicadores de costo, avance, beneficios y continuidad."),
            ("Bajo o medio poder / Alto interés", "Ciudadanos, arquitecto de software y equipo técnico", "Mantener informados e involucrados mediante pruebas, comunicación y canales de retroalimentación."),
            ("Bajo o medio poder / Bajo o medio interés", "Comité evaluador interno y entidades externas de biometría", "Monitorear y gestionar mediante contratos, acuerdos de servicio y reportes puntuales."),
        ],
    )

    add_paragraph_before(document, before, "Objetivos, plan de gestión, conflictos y riesgos", style="Heading 1")
    add_table_before(document, before, ["Interesado o grupo", "Objetivos relevantes", "Plan de gestión", "Conflictos y solución", "Beneficios", "Costos", "Riesgos y alternativas"], stakeholder_management)

    add_paragraph_before(document, before, "Beneficios y costos", style="Heading 1")
    add_table_before(document, before, ["Tipo", "Descripción"], benefit_cost_rows)

    add_paragraph_before(document, before, "Alternativas de gestión de riesgos", style="Heading 1")
    add_table_before(document, before, ["Riesgo", "Alternativa de gestión"], risk_alternatives)

    add_paragraph_before(document, before, "Conflictos derivados de la gestión de interesados", style="Heading 1")
    add_table_before(document, before, ["Conflicto", "Solución propuesta"], conflicts)

    normalize_alignment(document)
    clear_document_metadata(document)
    OUTPUT.mkdir(exist_ok=True)
    document.save(OUTPUT / "Matriz_poder_interes_diligenciada.docx")


if __name__ == "__main__":
    build_charter_from_template()
    build_matrix_from_template()
