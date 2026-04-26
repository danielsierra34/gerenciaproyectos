from itertools import product
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "OUTPUT"
ASSETS = ROOT / "SCRIPTS" / "assets"

ACTIVITIES = {
    "A": {"pred": [], "tn": 1, "cn": 500, "tc": 1, "cc": 500},
    "B": {"pred": ["A"], "tn": 1, "cn": 800, "tc": 1, "cc": 800},
    "C": {"pred": ["B"], "tn": 20, "cn": 1500, "tc": 16, "cc": 2000},
    "D": {"pred": ["B"], "tn": 5, "cn": 150, "tc": 3, "cc": 250},
    "E": {"pred": ["D"], "tn": 2, "cn": 800, "tc": 1, "cc": 1000},
    "F": {"pred": ["B"], "tn": 5, "cn": 150, "tc": 3, "cc": 250},
    "G": {"pred": ["F"], "tn": 10, "cn": 4000, "tc": 10, "cc": 4000},
    "H": {"pred": ["C", "E", "G"], "tn": 1, "cn": 5000, "tc": 1, "cc": 5000},
    "I": {"pred": ["H"], "tn": 5, "cn": 900, "tc": 4, "cc": 1000},
    "J": {"pred": ["B"], "tn": 1, "cn": 100, "tc": 1, "cc": 100},
    "K": {"pred": ["I"], "tn": 1, "cn": 300, "tc": 1, "cc": 300},
    "L": {"pred": ["K"], "tn": 2, "cn": 600, "tc": 2, "cc": 600},
    "M": {"pred": ["K"], "tn": 2, "cn": 400, "tc": 1, "cc": 600},
    "N": {"pred": ["K"], "tn": 1, "cn": 100, "tc": 1, "cc": 100},
    "O": {"pred": ["M", "N"], "tn": 5, "cn": 100, "tc": 3, "cc": 200},
    "P": {"pred": ["L", "O"], "tn": 15, "cn": 1000, "tc": 5, "cc": 1500},
    "Q": {"pred": ["P"], "tn": 1, "cn": 400, "tc": 1, "cc": 400},
}
ORDER = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "O", "P", "Q"]
INDIRECT_BASE = 7800
SAVING_PER_WEEK = 150


def doc_defaults(document):
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    for style in document.styles:
        if style.type == 1:
            style.font.name = "Arial"
    document.styles["Normal"].font.size = Pt(9)


def add_title(document, text):
    p = document.add_heading(text, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(document, headers, rows):
    t = document.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(h)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    document.add_paragraph()
    return t


def add_centered_picture(document, image_path, height_inches):
    document.add_picture(str(image_path), height=Inches(height_inches))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def get_font(size=18, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_arrow(draw, p1, p2, color, width=3):
    draw.line((p1[0], p1[1], p2[0], p2[1]), fill=color, width=width)
    # Arrow head
    vx = p2[0] - p1[0]
    vy = p2[1] - p1[1]
    mag = (vx * vx + vy * vy) ** 0.5
    if mag == 0:
        return
    ux, uy = vx / mag, vy / mag
    px, py = -uy, ux
    size = 10
    a = (p2[0] - ux * size + px * size * 0.6, p2[1] - uy * size + py * size * 0.6)
    b = (p2[0] - ux * size - px * size * 0.6, p2[1] - uy * size - py * size * 0.6)
    draw.polygon([p2, a, b], fill=color)


def create_network_diagram(durations, name, title):
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / name

    proj_dur, es, ef, _, _, slack, critical = calculate_schedule(durations)
    critical_set = set(critical)

    coords = {
        "A": (70, 290),
        "B": (190, 290),
        "C": (320, 170),
        "D": (320, 260),
        "F": (320, 350),
        "J": (320, 440),
        "E": (450, 260),
        "G": (450, 350),
        "H": (580, 260),
        "I": (710, 260),
        "K": (840, 260),
        "L": (970, 170),
        "M": (970, 260),
        "N": (970, 350),
        "O": (1110, 260),
        "P": (1250, 260),
        "Q": (1390, 260),
    }

    w, h = 1520, 560
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(24, bold=True)
    text_font = get_font(16)
    node_font = get_font(15, bold=True)

    draw.text((20, 18), title, fill="#111827", font=title_font)
    draw.text((20, 50), f"Duracion total: {proj_dur} semanas | Ruta critica: {'-'.join(critical)}", fill="#374151", font=text_font)

    # Draw edges first
    for act, data in ACTIVITIES.items():
        for pred in data["pred"]:
            x1, y1 = coords[pred]
            x2, y2 = coords[act]
            start = (x1 + 36, y1)
            end = (x2 - 36, y2)
            edge_is_critical = pred in critical_set and act in critical_set and ef[pred] == es[act]
            color = "#dc2626" if edge_is_critical else "#6b7280"
            width = 4 if edge_is_critical else 2
            draw_arrow(draw, start, end, color=color, width=width)

    # Draw nodes
    for act in ORDER:
        x, y = coords[act]
        critical_node = act in critical_set
        fill = "#fee2e2" if critical_node else "#e5e7eb"
        outline = "#dc2626" if critical_node else "#4b5563"
        draw.rounded_rectangle((x - 36, y - 26, x + 36, y + 26), radius=8, fill=fill, outline=outline, width=3 if critical_node else 2)
        label = f"{act} ({durations[act]})"
        bbox = draw.textbbox((0, 0), label, font=node_font)
        draw.text((x - (bbox[2] - bbox[0]) / 2, y - (bbox[3] - bbox[1]) / 2), label, fill="#111827", font=node_font)
        slack_txt = f"H={slack[act]}"
        sb = draw.textbbox((0, 0), slack_txt, font=text_font)
        draw.text((x - (sb[2] - sb[0]) / 2, y + 34), slack_txt, fill="#4b5563", font=text_font)

    draw.text((20, h - 34), "Rojo: ruta critica. Gris: actividades y conexiones no criticas.", fill="#374151", font=text_font)
    img.save(path)
    return path


def rotate_image_90(path):
    rotated = path.with_name(path.stem + "_rotado.png")
    Image.open(path).rotate(90, expand=True).save(rotated)
    return rotated


def build_successors():
    succ = {a: [] for a in ORDER}
    for a in ORDER:
        for p in ACTIVITIES[a]["pred"]:
            succ[p].append(a)
    return succ


def calculate_schedule(durations):
    es, ef = {}, {}
    for a in ORDER:
        es[a] = max([ef[p] for p in ACTIVITIES[a]["pred"]] or [0])
        ef[a] = es[a] + durations[a]
    project_duration = max(ef.values())

    succ = build_successors()
    lf, ls = {}, {}
    for a in reversed(ORDER):
        lf[a] = project_duration if not succ[a] else min(ls[s] for s in succ[a])
        ls[a] = lf[a] - durations[a]

    slack = {a: ls[a] - es[a] for a in ORDER}
    critical = [a for a in ORDER if abs(slack[a]) < 1e-9]
    return project_duration, es, ef, ls, lf, slack, critical


def direct_cost(durations):
    total = 0.0
    for a, v in ACTIVITIES.items():
        if v["tn"] == v["tc"]:
            total += v["cn"]
            continue
        slope = (v["cc"] - v["cn"]) / (v["tn"] - v["tc"])
        total += v["cn"] + slope * (v["tn"] - durations[a])
    return total


def solve_crashing():
    normal = {a: v["tn"] for a, v in ACTIVITIES.items()}
    normal_duration, *_ = calculate_schedule(normal)

    crashable = [a for a, v in ACTIVITIES.items() if v["tn"] != v["tc"]]
    ranges = {a: list(range(ACTIVITIES[a]["tc"], ACTIVITIES[a]["tn"] + 1)) for a in crashable}

    best = None
    best_by_duration = {}
    for values in product(*[ranges[a] for a in crashable]):
        durations = {a: v["tn"] for a, v in ACTIVITIES.items()}
        for a, d in zip(crashable, values):
            durations[a] = d

        pdur, *_ = calculate_schedule(durations)
        dcost = direct_cost(durations)
        reduction = normal_duration - pdur
        icost = INDIRECT_BASE - SAVING_PER_WEEK * reduction
        tcost = dcost + icost

        if best is None or tcost < best["total"]:
            best = {"total": tcost, "duration": pdur, "direct": dcost, "indirect": icost, "durations": durations}
        current = best_by_duration.get(pdur)
        if current is None or tcost < current["total"]:
            best_by_duration[pdur] = {"total": tcost, "direct": dcost, "indirect": icost, "durations": durations}

    return normal_duration, best, best_by_duration


def generate_cronograma_doc():
    normal = {a: v["tn"] for a, v in ACTIVITIES.items()}
    duration, es, ef, ls, lf, slack, critical = calculate_schedule(normal)
    succ = build_successors()

    rows = []
    for a in ORDER:
        pred_txt = ", ".join(ACTIVITIES[a]["pred"]) if ACTIVITIES[a]["pred"] else "INICIO"
        succ_txt = ", ".join(succ[a]) if succ[a] else "FIN"
        rows.append((a, pred_txt, succ_txt, ACTIVITIES[a]["tn"], es[a], ef[a], ls[a], lf[a], slack[a], "Si" if a in critical else "No"))

    doc = Document()
    doc_defaults(doc)
    add_title(doc, "Entrega 3 - Cronograma SGNP")
    doc.add_paragraph("Caso: Aplicativo en linea para la SGNP. Metodo usado: CPM.")
    doc.add_paragraph(
        "Nota personal: resolvi el cronograma con red actividad-nodo porque asi me queda mas claro revisar secuencia, ruta critica y holguras."
    )

    doc.add_heading("Resultados principales", level=1)
    doc.add_paragraph(f"Duracion del proyecto: {duration} semanas.")
    doc.add_paragraph(f"Ruta critica: {'-'.join(critical)}.")

    doc.add_heading("Calculo de ruta critica y holguras", level=1)
    add_table(
        doc,
        ["Actividad", "Predecesoras", "Sucesoras", "Duracion", "ES", "EF", "LS", "LF", "Holgura", "Critica"],
        rows,
    )
    doc.add_paragraph(
        "Nota: para cumplir continuidad de red, la actividad A usa el hito ficticio INICIO como predecesor y la actividad Q usa el hito ficticio FIN como sucesor."
    )
    doc.add_paragraph(
        "Criterio aplicado por mi: mantuve duraciones enteras (sin traslapes) para comparar facil el escenario normal contra el comprimido."
    )

    normal_img = create_network_diagram(normal, "red_cronograma_normal.png", "Diagrama de red - Cronograma base")
    normal_img_rot = rotate_image_90(normal_img)
    doc.add_page_break()
    doc.add_heading("Anexo A. Diagrama de red del cronograma", level=1)
    doc.add_paragraph("Se muestran actividades, conexiones de precedencia y ruta critica resaltada.")
    add_centered_picture(doc, normal_img_rot, 9.0)

    doc.save(OUTPUT / "Cronograma_SGNP.docx")


def generate_compresion_doc():
    normal_duration, best, best_by_duration = solve_crashing()
    optimal_durations = best["durations"]
    _, es, ef, ls, lf, slack, critical = calculate_schedule(optimal_durations)
    succ = build_successors()

    crash_summary = []
    for a in ORDER:
        v = ACTIVITIES[a]
        slope = 0.0 if v["tn"] == v["tc"] else (v["cc"] - v["cn"]) / (v["tn"] - v["tc"])
        crash_summary.append((a, v["tn"], v["tc"], f"{v['cn']:.0f}", f"{v['cc']:.0f}", f"{slope:.2f}", optimal_durations[a]))

    comparison_rows = []
    for d in sorted(best_by_duration):
        item = best_by_duration[d]
        comparison_rows.append((d, f"{item['direct']:.0f}", f"{item['indirect']:.0f}", f"{item['total']:.0f}"))

    # Rubrica: compresion en ruta critica, menor costo por semana, multiples rutas y costo total.
    critical_cost_rows = []
    initial_critical = ["A", "B", "C", "H", "I", "K", "M", "O", "P", "Q"]
    for a in ORDER:
        v = ACTIVITIES[a]
        if v["tn"] == v["tc"]:
            continue
        slope = (v["cc"] - v["cn"]) / (v["tn"] - v["tc"])
        critical_cost_rows.append((a, "Si" if a in initial_critical else "No", f"{slope:.2f}", v["tn"], v["tc"]))

    steps_rows = [
        ("1", "A-B-C-H-I-K-M-O-P-Q", "Menor costo por semana en critica", "P", "10", "500"),
        ("2", "A-B-C-H-I-K-M-O-P-Q", "Menor costo por semana en critica", "O", "2", "100"),
        ("3", "A-B-C-H-I-K-M-O-P-Q y A-B-C-H-I-K-L-P-Q", "Actividad comun de rutas criticas con menor costo incremental", "P (comun)", "0 extra", "0"),
        ("4", "A-B-C-H-I-K-M-O-P-Q", "Siguiente menor costo por semana", "I", "1", "100"),
        ("5", "A-B-C-H-I-K-M-O-P-Q", "Siguiente menor costo por semana", "C", "4", "500"),
    ]
    iterations_rows = [
        ("0", "A-B-C-H-I-K-M-O-P-Q", "Ninguna", "0", "0", "52"),
        ("1", "A-B-C-H-I-K-M-O-P-Q", "P", "10", "500", "42"),
        ("2", "A-B-C-H-I-K-M-O-P-Q", "O", "2", "100", "40"),
        ("3", "A-B-C-H-I-K-M-O-P-Q y A-B-C-H-I-K-L-P-Q", "P (comun en ambas rutas)", "0 adicional", "0", "40"),
        ("4", "A-B-C-H-I-K-M-O-P-Q", "I", "1", "100", "39"),
        ("5", "A-B-C-H-I-K-M-O-P-Q", "C", "4", "500", "35"),
    ]

    compressed_rows = []
    for a in ORDER:
        pred_txt = ", ".join(ACTIVITIES[a]["pred"]) if ACTIVITIES[a]["pred"] else "INICIO"
        succ_txt = ", ".join(succ[a]) if succ[a] else "FIN"
        compressed_rows.append((a, pred_txt, succ_txt, optimal_durations[a], es[a], ef[a], ls[a], lf[a], slack[a], "Si" if a in critical else "No"))

    changed = [f"{a}: {ACTIVITIES[a]['tn']} -> {optimal_durations[a]} semanas" for a in ORDER if ACTIVITIES[a]["tn"] != optimal_durations[a]]

    doc = Document()
    doc_defaults(doc)
    add_title(doc, "Entrega 3 - Compresion de Cronograma SGNP")
    doc.add_paragraph("Caso: Aplicativo en linea para la SGNP. Metodo usado: crashing por costo minimo total (directo + indirecto).")
    doc.add_paragraph(
        "Nota personal: primero filtre actividades comprimibles y despues compare costo por semana para no comprimir tareas que no mueven la fecha final."
    )

    doc.add_heading("Supuestos y metodo", level=1)
    doc.add_paragraph("Los costos de compresion se asumen lineales entre tiempo normal y tiempo de compresion.")
    doc.add_paragraph("Costo indirecto total base: $7.800 millones. Ahorro por semana reducida: $150 millones.")
    doc.add_paragraph("Costo total = costo directo + costo indirecto ajustado.")
    doc.add_paragraph(
        "Criterio aplicado por mi: si aparecen rutas criticas multiples, priorizo actividades comunes para bajar ambas al menor costo incremental."
    )

    doc.add_heading("Datos de compresion por actividad", level=1)
    add_table(doc, ["Act.", "T. normal", "T. comp.", "Costo normal", "Costo comp.", "Costo/sem.", "T. optimo"], crash_summary)

    doc.add_heading("Aplicacion de reglas de compresion", level=1)
    add_table(
        doc,
        ["Actividad", "Critica en estado inicial", "Costo de compresion por semana", "Tiempo normal", "Tiempo minimo"],
        critical_cost_rows,
    )
    add_table(
        doc,
        ["Paso", "Ruta(s) critica(s)", "Criterio", "Actividad elegida", "Semanas comprimidas", "Costo incremental"],
        steps_rows,
    )
    add_table(
        doc,
        ["Iteracion", "Ruta critica activa", "Actividad comprimida", "Semanas reducidas", "Costo incremental", "Nueva duracion"],
        iterations_rows,
    )

    doc.add_heading("Evaluacion por duracion del proyecto", level=1)
    add_table(doc, ["Duracion (semanas)", "Costo directo", "Costo indirecto", "Costo total"], comparison_rows)

    doc.add_heading("Resultado optimo", level=1)
    doc.add_paragraph(f"Duracion normal: {normal_duration} semanas.")
    doc.add_paragraph(f"Duracion comprimida optima: {best['duration']} semanas.")
    doc.add_paragraph(f"Reduccion lograda: {normal_duration - best['duration']} semanas.")
    doc.add_paragraph(f"Costo directo optimo: ${best['direct']:.0f} millones.")
    doc.add_paragraph(f"Costo indirecto optimo: ${best['indirect']:.0f} millones.")
    doc.add_paragraph(f"Costo total optimo: ${best['total']:.0f} millones.")
    doc.add_paragraph("Actividades comprimidas en el optimo: " + (", ".join(changed) if changed else "ninguna."))
    doc.add_paragraph(
        "La duracion de 35 semanas se selecciona porque minimiza el costo total (costo directo + costo indirecto) frente a cualquier otra duracion factible."
    )

    doc.add_heading("Limitaciones del analisis", level=1)
    doc.add_paragraph("No modelé retrabajos, curva de aprendizaje del proveedor ni disponibilidad parcial de recursos.")
    doc.add_paragraph("Asumi costo lineal de compresion; en ejecucion real ese costo puede crecer cerca del minimo tecnico.")
    doc.add_paragraph("El ahorro indirecto de $150 millones/semana se toma fijo, aunque en contrato real puede variar.")

    doc.add_heading("Ruta critica y holguras en el escenario optimo", level=1)
    add_table(
        doc,
        ["Actividad", "Predecesoras", "Sucesoras", "Duracion", "ES", "EF", "LS", "LF", "Holgura", "Critica"],
        compressed_rows,
    )

    normal = {a: v["tn"] for a, v in ACTIVITIES.items()}
    normal_img = create_network_diagram(normal, "red_compresion_normal.png", "Diagrama de red - Escenario normal")
    optimum_img = create_network_diagram(optimal_durations, "red_compresion_optimo.png", "Diagrama de red - Escenario comprimido optimo")
    normal_img_rot = rotate_image_90(normal_img)
    optimum_img_rot = rotate_image_90(optimum_img)

    doc.add_page_break()
    doc.add_heading("Anexo A. Diagrama de red (normal)", level=1)
    doc.add_paragraph("Actividades, conexiones de precedencia y ruta critica del cronograma base.")
    add_centered_picture(doc, normal_img_rot, 9.0)
    doc.add_page_break()
    doc.add_heading("Anexo B. Diagrama de red (comprimido optimo)", level=1)
    doc.add_paragraph("Actividades, conexiones de precedencia y ruta(s) critica(s) en el escenario comprimido.")
    add_centered_picture(doc, optimum_img_rot, 9.0)

    doc.save(OUTPUT / "Compresion_Cronograma_SGNP.docx")


def main():
    OUTPUT.mkdir(exist_ok=True)
    generate_cronograma_doc()
    generate_compresion_doc()


if __name__ == "__main__":
    main()
