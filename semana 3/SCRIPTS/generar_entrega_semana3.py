from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "OUTPUT"
ASSETS = ROOT / "SCRIPTS" / "assets"


WBS = [
    ("1.0", "Aplicativo en línea para trámites notariales SGNP", "Proyecto completo", "Acta de cierre del proyecto aprobada por Tinterillos y SGNP."),
    ("1.1", "Gerencia de proyectos", "Gestión integral del proyecto", "Plan, reportes, comités, control de cambios y cierre aprobados."),
    ("1.1.1", "Planeación del proyecto", "Plan de gestión, cronograma base, presupuesto base y matriz RACI.", "Plan aprobado por patrocinador y equipo ejecutor."),
    ("1.1.2", "Estimación y presupuesto", "Estimación de esfuerzo, costos, reservas y línea base de costos.", "Presupuesto base aprobado y trazable a entregables."),
    ("1.1.3", "Seguimiento y reportes", "Informes de avance, tablero de hitos, riesgos, costos y calidad.", "Reportes periódicos entregados y revisados en comité."),
    ("1.1.4", "Control de cambios", "Registro, análisis de impacto y decisiones sobre cambios de alcance.", "Solicitudes aprobadas, rechazadas o diferidas con soporte."),
    ("1.1.5", "Cierre del proyecto", "Acta de aceptación, lecciones aprendidas y transferencia operacional.", "Aceptación formal y documentación de cierre firmada."),
    ("1.2", "Diagnóstico y arquitectura", "Evaluación técnica y diseño objetivo de la solución.", "Arquitectura aprobada y brechas priorizadas."),
    ("1.2.1", "Diagnóstico de infraestructura", "Inventario de conectividad, servidores, ambientes, seguridad y protocolos.", "Informe de diagnóstico validado por entidades participantes."),
    ("1.2.2", "Dimensionamiento de demanda", "Volumen de datos, transacciones, concurrencia y crecimiento esperado.", "Modelo de capacidad aprobado con supuestos documentados."),
    ("1.2.3", "Arquitectura de integración", "Diseño de microservicios, integraciones y comunicaciones con entidades.", "Documento de arquitectura aprobado por Turing y Tinterillos."),
    ("1.2.4", "Diseño de seguridad", "Controles de confidencialidad, integridad, autenticación y auditoría.", "Controles trazados a riesgos y requisitos de identidad."),
    ("1.3", "Funcionalidades notariales en línea", "Módulos funcionales para compraventa de inmuebles.", "Funcionalidades desarrolladas y probadas en ambiente de pruebas."),
    ("1.3.1", "Citas virtuales", "Programación y administración de citas entre partes y notaría.", "Citas creadas, modificadas y consultadas correctamente."),
    ("1.3.2", "Identificación biométrica", "Validación facial, por voz y por huella de participantes.", "Identidad validada con tres factores y registro de auditoría."),
    ("1.3.3", "Gestión documental", "Carga, consulta y control de escrituras electrónicas.", "Escrituras disponibles, versionadas y protegidas."),
    ("1.3.4", "Anotaciones documentales", "Comentarios y observaciones de participantes sobre documentos.", "Anotaciones registradas y visibles para revisión notarial."),
    ("1.3.5", "Flujos de trabajo", "Automatización del proceso de compraventa de inmuebles.", "Estados del trámite ejecutados según reglas definidas."),
    ("1.4", "Plataforma tecnológica", "Infraestructura, ambientes y operación técnica.", "Ambientes listos para pruebas, piloto y producción."),
    ("1.4.1", "Ambientes de desarrollo y pruebas", "Configuración de ambientes, datos de prueba y despliegues.", "Ambientes disponibles y documentados."),
    ("1.4.2", "Escalabilidad y desempeño", "Capacidad, balanceo, tiempos de respuesta y pruebas de carga.", "Resultados dentro de umbrales aprobados."),
    ("1.4.3", "Disponibilidad y continuidad", "Monitoreo, respaldo, recuperación y contingencia.", "Plan probado con evidencia de recuperación."),
    ("1.4.4", "Integración con plataforma actual", "Conexión con microservicios Spring Boot y servicios existentes.", "Integraciones funcionando con pruebas end-to-end."),
    ("1.5", "Calidad, pruebas y aceptación", "Verificación funcional, técnica y de seguridad.", "Evidencia de pruebas y aceptación del usuario."),
    ("1.5.1", "Pruebas funcionales", "Casos de prueba de citas, documentos, biometría y flujos.", "Casos críticos ejecutados con defectos cerrados."),
    ("1.5.2", "Pruebas de seguridad", "Validación de acceso, identidad, auditoría e integridad documental.", "Hallazgos críticos cerrados antes de producción."),
    ("1.5.3", "Pruebas de desempeño", "Carga, concurrencia y tiempos de respuesta.", "Resultados cumplen criterios definidos."),
    ("1.5.4", "Aceptación de usuarios", "Validación con notarios, funcionarios y representantes de usuarios.", "Acta de aceptación de piloto firmada."),
    ("1.6", "Gestión del cambio y despliegue", "Preparación organizacional y salida a producción.", "Usuarios preparados y solución operando en notarías objetivo."),
    ("1.6.1", "Capacitación", "Materiales, sesiones y guías para notarios y funcionarios.", "Usuarios clave capacitados y asistencia registrada."),
    ("1.6.2", "Comunicaciones", "Mensajes de adopción, instrucciones y canales de soporte.", "Comunicaciones publicadas a públicos definidos."),
    ("1.6.3", "Piloto controlado", "Operación limitada en notarías seleccionadas.", "Piloto evaluado con incidencias priorizadas."),
    ("1.6.4", "Puesta en producción", "Despliegue final y estabilización inicial.", "Funcionalidades en operación antes del cierre del primer año."),
]


COST_FUNCTIONALITY = 142_500_000
COST_EXECUTION_BASE = 210_000_000
COST_BREAKDOWN = [
    ("1.1.1", "Fase de inicio", 0.03, "Inicio formal, coordinación del proveedor y preparación del piloto."),
    ("1.1.2", "Definición de producto", 0.10, "Alcance funcional de anotaciones documentales, criterios de aceptación y backlog del piloto."),
    ("1.2.3", "Definición y actualización de arquitectura", 0.08, "Arquitectura del piloto alineada con microservicios y servicios de la notaría."),
    ("1.2.4", "Definición e implementación de seguridad", 0.16, "Seguridad equivalente al proyecto real, identidad, trazabilidad y controles de acceso."),
    ("1.4.1", "Infraestructura tecnológica para desarrollo y pruebas", 0.07, "Ambientes para construcción, pruebas y soporte de 5.000 usuarios concurrentes."),
    ("1.3.4", "Desarrollo de software", 0.42, "Funcionalidad de anotaciones documentales y componentes de integración necesarios."),
    ("1.5.3", "Pruebas de carga y estrés", 0.03, "Validación de concurrencia, desempeño y estabilidad del piloto."),
    ("1.5.4", "Pruebas piloto del software", 0.03, "Validación operativa del piloto con usuarios y escenarios representativos."),
    ("1.6.4", "Mantenimiento y soporte", 0.08, "Soporte durante estabilización y atención de incidentes del piloto."),
]


def doc_defaults(document):
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    for style in document.styles:
        if style.type == 1:
            style.font.name = "Arial"
    document.styles["Normal"].font.size = Pt(9)


def add_title(document, text):
    paragraph = document.add_heading(text, 0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    document.add_paragraph()
    return table


def money(value):
    return f"${value:,.0f}".replace(",", ".")


def get_font(size=22, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw, xy, title, fill, outline="#1f2937"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=2)
    font = get_font(21, bold=True)
    small = get_font(18)
    title_lines = wrap_text(draw, title, font, x2 - x1 - 18)
    y = y1 + 12
    for line in title_lines[:2]:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#111827", font=font)
        y += 25
    if len(title_lines) > 2:
        draw.text((x1 + 10, y), " ".join(title_lines[2:]), fill="#111827", font=small)


def create_wbs_tree_image():
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / "edt_arbol.png"
    img = Image.new("RGB", (1700, 980), "#ffffff")
    draw = ImageDraw.Draw(img)

    root = (460, 35, 1240, 130)
    level2 = [
        ("1.1 Gerencia de\nproyectos", (20, 220, 280, 320)),
        ("1.2 Diagnóstico y\narquitectura", (300, 220, 560, 320)),
        ("1.3 Funcionalidades\nnotariales en línea", (580, 220, 840, 320)),
        ("1.4 Plataforma\ntecnológica", (860, 220, 1120, 320)),
        ("1.5 Calidad, pruebas\ny aceptación", (1140, 220, 1400, 320)),
        ("1.6 Gestión del cambio\ny despliegue", (1420, 220, 1680, 320)),
    ]
    children = {
        "1.1 Gerencia de\nproyectos": ["1.1.1 Planeación", "1.1.2 Estimación", "1.1.3 Reportes", "1.1.4 Control de\ncambios", "1.1.5 Cierre"],
        "1.2 Diagnóstico y\narquitectura": ["1.2.1 Infraestructura", "1.2.2 Demanda", "1.2.3 Integración", "1.2.4 Seguridad"],
        "1.3 Funcionalidades\nnotariales en línea": ["1.3.1 Citas\nvirtuales", "1.3.2 Identificación\nbiométrica", "1.3.3 Gestión\ndocumental", "1.3.4 Anotaciones", "1.3.5 Flujos"],
        "1.4 Plataforma\ntecnológica": ["1.4.1 Ambientes", "1.4.2 Escalabilidad", "1.4.3 Continuidad", "1.4.4 Integración\nactual"],
        "1.5 Calidad, pruebas\ny aceptación": ["1.5.1 Funcionales", "1.5.2 Seguridad", "1.5.3 Desempeño", "1.5.4 Aceptación"],
        "1.6 Gestión del cambio\ny despliegue": ["1.6.1 Capacitación", "1.6.2 Comunicaciones", "1.6.3 Piloto", "1.6.4 Producción"],
    }

    draw_box(draw, root, "1.0 Aplicativo en línea para trámites notariales SGNP", "#d9ead3")
    root_bottom = ((root[0] + root[2]) // 2, root[3])
    trunk_y = 175
    draw.line((root_bottom[0], root_bottom[1], root_bottom[0], trunk_y), fill="#111827", width=3)
    draw.line((150, trunk_y, 1550, trunk_y), fill="#111827", width=3)

    for title, box in level2:
        center = ((box[0] + box[2]) // 2, box[1])
        draw.line((center[0], trunk_y, center[0], center[1]), fill="#111827", width=3)
        draw_box(draw, box, title, "#cfe2f3")
        child_items = children[title]
        y = 420
        for item in child_items:
            child_box = (box[0], y, box[2], y + 72)
            draw.line(((box[0] + box[2]) // 2, box[3], (box[0] + box[2]) // 2, child_box[1]), fill="#6b7280", width=2)
            draw_box(draw, child_box, item, "#f3f4f6", "#6b7280")
            y += 90

    img.save(path)
    return path


def build_wbs_doc():
    document = Document()
    doc_defaults(document)
    add_title(document, "Estructura de Desglose del Trabajo (EDT)")
    document.add_paragraph(
        "Proyecto: Aplicativo en línea para la SGNP. La EDT está orientada a entregables, usa sustantivos y permite verificar cuándo cada componente está completo."
    )
    document.add_heading("EDT en formato árbol", level=1)
    document.add_paragraph(
        "El diagrama en árbol se presenta en el Anexo A al final del documento, en versión rotada para ajuste a hoja."
    )
    document.add_heading("EDT en formato tabular", level=1)
    add_table(
        document,
        ["Código", "Entregable", "Descripción", "Criterio de finalización"],
        WBS,
    )

    tree_path = create_wbs_tree_image()
    rotated_path = ASSETS / "edt_arbol_rotado.png"
    Image.open(tree_path).rotate(90, expand=True).save(rotated_path)

    document.add_page_break()
    document.add_heading("Anexo A. EDT en árbol (rotado 90°)", level=1)
    document.add_picture(str(rotated_path), width=Inches(5.1))
    document.save(OUTPUT / "EDT_SGNP.docx")


def build_cost_doc():
    base_rows = []
    for code, item, pct, note in COST_BREAKDOWN:
        base_rows.append((code, item, f"{pct:.0%}", money(COST_EXECUTION_BASE * pct), note))

    base_pilot = COST_FUNCTIONALITY + COST_EXECUTION_BASE

    modified_breakdown = []
    modified_execution = 0
    for code, item, pct, note in COST_BREAKDOWN:
        new_pct = pct
        if item == "Desarrollo de software":
            new_pct = 0.62
        if item in ("Pruebas de carga y estrés", "Pruebas piloto del software"):
            new_pct = 0
        amount = COST_EXECUTION_BASE * new_pct
        modified_execution += amount
        modified_breakdown.append((code, item, f"{new_pct:.0%}", money(amount), note))
    modified_pilot = COST_FUNCTIONALITY + modified_execution

    infrastructure_base = COST_EXECUTION_BASE * 0.07
    load_tests_base = COST_EXECUTION_BASE * 0.03
    scaled_infrastructure = infrastructure_base * (1_000 / 5_000)
    scaled_load_tests = load_tests_base * (1_000 / 5_000)
    user_reduction_saving = (infrastructure_base - scaled_infrastructure) + (load_tests_base - scaled_load_tests)
    adjusted_execution_1000 = COST_EXECUTION_BASE - user_reduction_saving
    adjusted_pilot_1000 = COST_FUNCTIONALITY + adjusted_execution_1000

    document = Document()
    doc_defaults(document)
    add_title(document, "Estimación de Costos")
    document.add_paragraph("Proyecto: Aplicativo en línea para la SGNP. Alcance de estimación: piloto de la funcionalidad de anotaciones documentales.")

    document.add_heading("Supuestos de la estimación", level=1)
    assumptions = [
        "El costo de la funcionalidad de anotaciones documentales informado en el enunciado es $142.500.000.",
        "El costo de ejecución del piloto se estima por analogía con proyectos anteriores de Turing cuyo costo fue $210.000.000.",
        "El piloto debe conservar características del proyecto real: disponibilidad de 99,99%, seguridad, desempeño y facilidad de uso.",
        "El piloto base se dimensiona para 5.000 usuarios concurrentes.",
        "El personal interno de Turing está ocupado, por lo que el desarrollo y las pruebas del piloto se contratan con un proveedor.",
        "Los porcentajes históricos se aplican sobre el costo de ejecución del piloto, no sobre el costo de la funcionalidad.",
        "Los valores están expresados en pesos colombianos y no incluyen impuestos, retenciones ni cambios de alcance posteriores.",
        "La EDT se usa para relacionar los rubros con entregables: producto, arquitectura, seguridad, infraestructura, desarrollo, pruebas, soporte y gerencia.",
    ]
    for item in assumptions:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Nivel de precisión y exactitud", level=1)
    add_table(
        document,
        ["Elemento", "Definición aplicada"],
        [
            ("Nivel de precisión", "Los valores se presentan en pesos colombianos y se redondean al peso. Para presentación ejecutiva se pueden redondear a millones de pesos."),
            ("Exactitud esperada", "Estimación de orden de magnitud con rango esperado de -25% a +75%, porque se basa en información histórica y aún no en cotizaciones detalladas del proveedor."),
            ("Unidad de medida", "Porcentaje histórico aplicado al costo de ejecución del piloto."),
            ("Moneda", "Pesos colombianos (COP)."),
        ],
    )

    document.add_heading("Método de estimación", level=1)
    document.add_paragraph(
        "Se usa estimación análoga de orden de magnitud. El costo de ejecución del piloto se toma de un proyecto anterior de Turing y se distribuye con los porcentajes históricos entregados en el enunciado. "
        "El costo total del piloto se calcula como: Costo del Piloto = Costo de la Funcionalidad + Costo de Ejecución del Piloto."
    )

    document.add_heading("1. Costo estimado del piloto en orden de magnitud", level=1)
    add_table(
        document,
        ["Código EDT", "Rubro", "Porcentaje histórico", "Costo estimado", "Justificación"],
        base_rows,
    )
    add_table(
        document,
        ["Concepto", "Valor"],
        [
            ("Costo de la funcionalidad", money(COST_FUNCTIONALITY)),
            ("Costo de ejecución del piloto", money(COST_EXECUTION_BASE)),
            ("Costo estimado del piloto", money(base_pilot)),
        ],
    )

    document.add_heading("2. Cambio si desarrollo sube a 62% y se eliminan pruebas de carga, estrés y piloto", level=1)
    document.add_paragraph(
        "El desarrollo de software aumenta de 42% a 62%. Las pruebas de carga y estrés y las pruebas piloto se eliminan, por lo que cada una pasa de 3% a 0%. "
        "El efecto neto sobre la ejecución es un aumento de 14 puntos porcentuales: +20 puntos por desarrollo y -6 puntos por las pruebas eliminadas."
    )
    add_table(
        document,
        ["Código EDT", "Rubro", "Nuevo porcentaje", "Nuevo costo", "Comentario"],
        modified_breakdown,
    )
    add_table(
        document,
        ["Concepto", "Valor"],
        [
            ("Costo de ejecución modificado", money(modified_execution)),
            ("Costo de la funcionalidad", money(COST_FUNCTIONALITY)),
            ("Nuevo costo estimado del piloto", money(modified_pilot)),
            ("Incremento frente al escenario base", money(modified_pilot - base_pilot)),
        ],
    )

    document.add_heading("3. Mejora de la estimación con 1.000 usuarios concurrentes", level=1)
    document.add_paragraph(
        "La reducción de 5.000 a 1.000 usuarios concurrentes permite mejorar la estimación separando costos fijos de software y costos variables de capacidad. "
        "No conviene dividir todo el costo entre cinco, porque producto, arquitectura, seguridad, desarrollo y soporte siguen siendo necesarios. "
        "Como aproximación paramétrica, se escala al 20% únicamente la infraestructura tecnológica y las pruebas de carga y estrés, que son los rubros más sensibles a concurrencia."
    )
    add_table(
        document,
        ["Rubro sensible a usuarios", "Costo base 5.000 usuarios", "Factor 1.000/5.000", "Costo ajustado", "Ahorro estimado"],
        [
            ("Infraestructura tecnológica para desarrollo y pruebas", money(infrastructure_base), "20%", money(scaled_infrastructure), money(infrastructure_base - scaled_infrastructure)),
            ("Pruebas de carga y estrés", money(load_tests_base), "20%", money(scaled_load_tests), money(load_tests_base - scaled_load_tests)),
            ("Total ajuste", "", "", "", money(user_reduction_saving)),
        ],
    )
    add_table(
        document,
        ["Concepto", "Valor"],
        [
            ("Costo de ejecución ajustado para 1.000 usuarios", money(adjusted_execution_1000)),
            ("Costo de la funcionalidad", money(COST_FUNCTIONALITY)),
            ("Costo estimado del piloto ajustado", money(adjusted_pilot_1000)),
        ],
    )

    document.add_heading("Resultado", level=1)
    document.add_paragraph(
        f"El costo estimado del piloto base es {money(base_pilot)}. Si el desarrollo sube a 62% y se eliminan las pruebas indicadas, el costo sube a {money(modified_pilot)}. "
        f"Si la concurrencia baja a 1.000 usuarios y se ajustan solo los rubros sensibles a capacidad, el costo estimado baja a {money(adjusted_pilot_1000)}. "
        "Para mejorar la exactitud, Turing debería solicitar cotización formal del proveedor, separar costos fijos y variables, validar precios de nube/ambientes, y conservar pruebas mínimas de carga y piloto para no elevar el riesgo operativo."
    )

    document.save(OUTPUT / "Estimacion_Costos_SGNP.docx")


def main():
    OUTPUT.mkdir(exist_ok=True)
    build_wbs_doc()
    build_cost_doc()


if __name__ == "__main__":
    main()
